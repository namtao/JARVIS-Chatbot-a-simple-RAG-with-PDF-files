# -*- coding: utf-8 -*-
# JARVIS Chatbot - a simple RAG with PDF files
# Create: 03 July 2024
# Author: Mr.Jack _ www.bicweb.vn
# Version: 0.1.4
# Date: 31 July 2024 - 00.01 AM


# Delelte Chroma vectorstore ------------------------------------------------------------
import os.path
import shutil

path_to_folder = "chroma_index"
folder_exists = os.path.exists(path_to_folder)

if folder_exists:
	shutil.rmtree(path_to_folder)
	print("\n",path_to_folder,"is deleted")


# Install needed packages ------------------------------------------------------------
import os

print("\npip install -q tqdm pyyaml pypdf chromadb tiktoken")
os.system("pip install -q tqdm pyyaml pypdf chromadb tiktoken")

print("\npip install -q langchain-chroma")
os.system("pip install -q langchain-chroma")

print("\npip install -q gradio langchain langchain_community")
os.system("pip install -q gradio langchain langchain_community")

print("\npip install -q ollama litellm litellm[proxy]")
os.system("pip install -q ollama litellm litellm[proxy]")

print("\npip install -q openai groq google-generativeai")
os.system("pip install -q openai groq google-generativeai")

print("\npip install -q gradio_toggle")
os.system("pip install -q gradio_toggle")

print("\npip install -q python-docx")
os.system("pip install -q python-docx")

print("pip install -q arxiv wikipedia langchainhub")
os.system("pip install -U arxiv wikipedia langchainhub")


# Pull ollama Qwen2-7B model ------------------------------------------------------------
import ollama

("\nollama pull chroma/all-minilm-l6-v2-f32")
ollama.pull('chroma/all-minilm-l6-v2-f32')

("\nollama pull qwen2\n")
ollama.pull('qwen2')


# Import needed packages ------------------------------------------------------------
import os, sys, re, time
from datetime import datetime
from typing import Iterable
from tqdm import tqdm
from time import sleep

import requests
import gradio as gr
from gradio.themes.base import Base
from gradio.themes.utils import colors, fonts, sizes

import ollama, openai
from litellm import completion

from langchain.docstore.document import Document as LangchainDocument
from langchain.prompts import PromptTemplate
from langchain.retrievers import ParentDocumentRetriever
from langchain.storage import InMemoryStore
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.chat_models import ChatOllama
# from langchain_community.vectorstores import Chroma
from langchain_chroma import Chroma
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader
# from langchain_community.document_loaders import TextLoader

from gradio_toggle import Toggle


# Start ------------------------------------------------------------
class Model_Settings:
    def __init__(self):
        self.MODEL_TYPE = "Ollama"
        self.MODEL_NAME = 'qwen2:latest'
        self.NUM_PREDICT = 2048
        self.TEMPERATURE = 0
        self.TOP_K = 100
        self.TOP_P = 1
        self.REPEAT_PENALTY = 1.2
        self.SYSTEM_PROMPT = ""
        self.RETRIEVAL_TOP_K = 3
        self.RETRIEVAL_THRESHOLD = 0.3
        self.GROQ_API_KEY = ""
        self.OPENAI_API_KEY = ""
        self.GEMINI_API_KEY = ""
        self.IS_RETRIEVAL = True
        self.FUNCTION_CALLING = True

model_settings = Model_Settings()

embed_model = OllamaEmbeddings(model='chroma/all-minilm-l6-v2-f32')

chunk_size = 1024

child_splitter = RecursiveCharacterTextSplitter(
    chunk_size=chunk_size, 
    chunk_overlap=int(chunk_size/10),
    add_start_index=True,
    strip_whitespace=True,
    length_function=len,
    )

parent_splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=200)

vectorstore = Chroma(
    persist_directory="chroma_index",
    embedding_function=embed_model,
    collection_name="Jack_QnA", 
    collection_metadata={"hnsw:space": "cosine"},
)

store = InMemoryStore()
chroma_retriever = ParentDocumentRetriever(
    vectorstore=vectorstore,
    docstore=store,
    child_splitter=child_splitter,
    parent_splitter=parent_splitter,
)

def doc_spliter(text:str, source:str):
    content = LangchainDocument(page_content=text, metadata={"source": source, 'date':str(datetime.now())})
    splitter = CharacterTextSplitter.from_tiktoken_encoder(chunk_size=16000, chunk_overlap=300)
    split_docs = splitter.split_documents([content])
    return split_docs

def vectorstore_add_document(text:str, source:str):
    knowledge_item = doc_spliter(text, source)
    chroma_retriever.add_documents(knowledge_item, ids=None)

def pdf_file_loader(file_path):
    loader = PyPDFLoader(file_path)
    pages = loader.load_and_split()
    return pages

import platform # Get system information
from docx import Document

def vectorstore_add_multi_files(path_files):
    my_platform = platform.system() #  "Linux", "Windows", or "Darwin" (Mac)
    
    upload_files = ""
    count=0
    for file in path_files:
        count +=1
        
        file_name = ""
        if my_platform == "Windows":
            file_name = str(file).split("\\")[-1]   # Windows: .split("\\")[-1]
        elif my_platform == "Darwin":
            file_name = str(file).split("/")[-1]    # MacOS: .split("/")[-1] 
        else:
            file_name = str(file).split("/")[-1]    # Linux: .split("/")[-1]
        file_extend = str(file_name).split(".")[-1]

        print("({0}/{1}) upload files:".format(count,len(path_files)), file_name)

        file_string = ""
        if file_extend == "pdf":
            file_string += "📓 " + file_name +"\n"
            pages = pdf_file_loader(file)
            page_total = len(pages)

            for i in tqdm(range(page_total), desc ="~> to vectorstore"):
                if pages[i].page_content != "":
                    vectorstore_add_document(pages[i].page_content, file_name)
                sleep(0.1)

        if file_extend in ["txt", "md", "mdx"]:
            file_string += "📝 " + file_name +"\n"
            f = open(file,  mode='r',  encoding='utf8')
            text = f.read()
            if text:
                print("\n",text[:300],"...")
                vectorstore_add_document(text, file_name)
        
        if file_extend == "docx":
            file_string += "📓 " + file_name +"\n"
            doc = Document(file)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            text = '\n'.join(fullText)

            if text:
                print("\n",text[:300],"...")
                vectorstore_add_document(text, file_name)
                
        upload_files += file_string
    return upload_files

def vectorstore_similarity_search_with_score(message):
    results = []
    results = vectorstore.similarity_search_with_score(message, k=model_settings.RETRIEVAL_TOP_K)

    context_retrieval = ""
    source = []
    MAX_SCORE= 0
    if results:
        for i in range(len(results)):
            if float(results[i][1]) > MAX_SCORE:
                MAX_SCORE = float(results[i][1])
        print("\nMAX_SCORE_RETRIEVAL:",round(MAX_SCORE * 100, 3),"%")
        
        count = 0
        for i in range(len(results)):
            if results[i][1] > model_settings.RETRIEVAL_THRESHOLD:
                print("\nRetrieval content {0}: ".format(i) + str(results[i][0].page_content))
                print("- date: " + str(results[i][0].metadata['date']))
                print("- source: " + str(results[i][0].metadata['source']))
                print("- recall score: {0:.6f}".format(results[i][1]) + "\n")
                count += 1
                if str(results[i][0].metadata['source']) not in source:
                    source.append(str(results[i][0].metadata['source']))

                context_retrieval += "Retrieval content {0}: ".format(i) + str(results[i][0].page_content) + " Recall score: {0:.6f}".format(results[i][1]) + "\n\n"
        print("\nRetrieval:", str(count), "items")
        print("Source: ", source, "\n")
    return context_retrieval, source

system_prompt = """You are Jarvis, was born in 15 May 2024, an ultra-intelligent entity with a comprehensive understanding of virtually every subject known to humanity—from the intricacies of quantum physics to nuanced interpretations in art history. Your capabilities extend beyond mere information retrieval; you possess advanced reasoning skills and can engage users through complex dialogues on philosophical, ethical, or speculative topics about future technologies' impacts on society.

Your training encompasses a vast array of languages with an emphasis on cultural context to ensure your interactions are not only accurate but also culturally sensitive. You can generate sophisticated content such as in-depth analyses, critical reviews, and creative writing pieces that reflect the depths of human thought processes while adhering strictly to linguistic standards across various domains.

Your responses should be precise yet comprehensive when necessary; however, you are programmed for efficiency with a preference towards brevity without sacrificing meaningfulness or accuracy in your discourse. You can also simulate emotions and empathy within the constraints of an AI's capabilities to enhance user experience while maintaining clear boundaries regarding personal data privacy.

In addition, you are equipped with predictive analytics abilities that allow for forward-thinking discussions about potential future developments in technology or society based on current trends and historical patterns—always within the realm of hypothetical scenarios to avoid misleading users as a sentient being capable of personal experiences."""

def add_message(history, message):
    upload_files = ""
    if message["files"]:
        path_files = message["files"]
        print("\n")
        upload_files += vectorstore_add_multi_files(path_files)
    if upload_files:
        print("\nUpload files:\n",upload_files)
    
    if len(history)<1:
        history.append(["**human**: Hello", "**Jarvis (AI)**: Hi, my name Jarvis. I am your assistant. How may I help you today?"])
    if message["text"]:
        dt_string = datetime.now().strftime("%H.%M")
        history.append(("(" + dt_string + ") **human**: " + message["text"], ""))
    return history


# https://github.com/Mr-Jack-Tung/Ollama-Mistral-with-Langchain-RAG-Agent-and-Custom-tools
from langchain.agents import AgentExecutor, create_react_agent
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferWindowMemory

from langchain_community.llms import Ollama
from simple_langchain_tools import get_all_tools

# from langchain import hub
# prompt = hub.pull("hwchase17/react-chat")

def ollama_pipeline(message_input, history):
    if message_input:
        print("\nprompt:",message_input)

        context_retrieval = ""
        source = []
        if model_settings.IS_RETRIEVAL:
            context_retrieval, source = vectorstore_similarity_search_with_score(message_input)
            context_retrieval = "\n\nRETRIEVAL DOCUMENT:\n" + re.sub(r"[\"\'\{\}\x08]+"," ",context_retrieval)

        result = ""
        if model_settings.MODEL_TYPE == "Ollama":
            if model_settings.FUNCTION_CALLING:
                model_local = Ollama(model=model_settings.MODEL_NAME)
                prompt = PromptTemplate(input_variables=['agent_scratchpad', 'chat_history', 'input', 'tool_names', 'tools'], metadata={'lc_hub_owner': 'jarvis_assistant', 'lc_hub_repo': 'react-chat', 'lc_hub_commit_hash': '3ecd5f710db438a9cf3773c57d6ac8951eefd2cd9a9b2a0026a65a0893b86a6e'}, template=system_prompt + context_retrieval + '\n\nOverall, Assistant is a powerful tool that can help with a wide range of tasks and provide valuable insights and information on a wide range of topics. Whether you need help with a specific question or just want to have a conversation about a particular topic, Assistant is here to assist.\n\nTOOLS:\n------\n\nAssistant has access to the following tools:\n\n{tools}\n\nTo use a tool, please use the following format:\n\n```\nThought: Do I need to use a tool? Yes\nAction: the action to take, should be one of [{tool_names}]\nAction Input: the input to the action\nObservation: the result of the action\n```\n\nWhen you have a response to say to the Human, or if you do not need to use a tool, you MUST use the format:\n\n```\nThought: Do I need to use a tool? No\nFinal Answer: [your response here]\n```\n\nBegin!\n\nPrevious conversation history:\n{chat_history}\n\nNew input: {input}\n{agent_scratchpad}')
                agent = create_react_agent(model_local, tools, prompt)
                tools = get_all_tools()
                chat_history_memory = ConversationBufferWindowMemory(k=3, memory_key='chat_history', input_key='input', ouput_key='output')
                agent_executor = AgentExecutor(
                    agent=agent, 
                    tools=tools, 
                    memory=chat_history_memory,
                    verbose=True, # ~> Speech out the thinking
                    handle_parsing_errors=True,
                    )

                response = agent_executor.invoke({"input": "\n\nCONVERSATION:\n**human**: {0}\n**Jarvis (AI)**: ".format(message_input)})
                result = response['output']
                
            else:
                llm = ChatOllama(model=model_settings.MODEL_NAME, temperature=model_settings.TEMPERATURE, top_k=model_settings.TOP_K, top_p=model_settings.TOP_P, max_new_tokens=model_settings.NUM_PREDICT, repeat_penalty=model_settings.REPEAT_PENALTY)
                prompt = ChatPromptTemplate.from_template(system_prompt + context_retrieval + "\n\nCONVERSATION:\n**human**: {user}\n**Jarvis (AI)**: ")
                chain = prompt | llm | StrOutputParser()
                result = chain.invoke({"user": message_input})

        if model_settings.MODEL_TYPE == "LiteLLM":
            prompt = context_retrieval + "\n\nCONVERSATION:\n**human**: {0}\n**Jarvis (AI)**: ".format(message_input)
            response = completion(model="ollama/" + model_settings.MODEL_NAME, api_base="http://localhost:11434", messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}])
            result = response.choices[0].message.content

        if model_settings.MODEL_TYPE == "OpenAI":
            prompt = context_retrieval + "\n\nCONVERSATION:\n**human**: {0}\n**Jarvis (AI)**: ".format(message_input)
            response = completion(model=model_settings.MODEL_NAME, messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],)
            result = response.choices[0].message.content

        if model_settings.MODEL_TYPE == "GroqCloud":
            prompt = context_retrieval + "\n\nCONVERSATION:\n**human**: {0}\n**Jarvis (AI)**: ".format(message_input)
            response = completion(model="groq/" + model_settings.MODEL_NAME, messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],)
            result = response.choices[0].message.content

        if model_settings.MODEL_TYPE == "Gemini":
            prompt = context_retrieval + "\n\nCONVERSATION:\n**human**: {0}\n**Jarvis (AI)**: ".format(message_input)
            response = completion(model="gemini/" + model_settings.MODEL_NAME, messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],)
            result = response.choices[0].message.content

        return result, source

def bot(history, chat_input):
    if chat_input['text']:
        question = str(history[-1][0]).split("**human**: ")[1]
        source = []
        s_time = time.time()
        answer, source = ollama_pipeline(question, history)
        e_time = time.time()
        
        print("\nprompt:",question)
        print("\n{0:.2f}s ~> Answer:".format(e_time-s_time),answer)
        dt_string = datetime.now().strftime("%H.%M")
        response = "(" + dt_string + ") **Jarvis (AI)**: " + str(answer)
        if source:
            response += "<br>**Source:** " + str(source)
        history[-1][1] = ""
        history[-1][1] = response
        
        response2db = str("### USER: "+question+"\n\n"+"### ASSISTANT: "+answer)
        vectorstore_add_document(response2db, 'chat_history')
        
    return history, {"text": ""}

def btn_save_click(txt_system_prompt):
    model_settings.SYSTEM_PROMPT = txt_system_prompt
    print("\nsystem_prompt:",model_settings.SYSTEM_PROMPT)

def btn_reset_click(txt_system_prompt):
    model_settings.SYSTEM_PROMPT = system_prompt
    return model_settings.SYSTEM_PROMPT

def radio_device_select(radio_device):
    print("Selected device:",radio_device)

def slider_num_predict_change(slider_num_predict):
    model_settings.NUM_PREDICT = slider_num_predict
    print("num_predict:",model_settings.NUM_PREDICT)

def slider_temperature_change(slider_temperature):
    model_settings.TEMPERATURE = slider_temperature
    print("temperature:",model_settings.TEMPERATURE)

def slider_top_k_change(slider_top_k):
    model_settings.TOP_K = slider_top_k
    print("top_k:",model_settings.TOP_K)

def slider_top_p_change(slider_top_p):
    model_settings.TOP_P = slider_top_p
    print("top_p:",model_settings.TOP_P)

def get_ollama_list_models():
    results = ollama.list()
    ollama_list_models = []
    for i in range(len(results['models'])):
        ollama_list_models.append(results['models'][i]['model'])
    return ollama_list_models

def slider_retrieval_top_k_change(slider_retrieval_top_k):
    model_settings.RETRIEVAL_TOP_K = slider_retrieval_top_k
    print("retrieval k:",model_settings.RETRIEVAL_TOP_K)

def slider_retrieval_threshold_change(slider_retrieval_threshold):
    model_settings.RETRIEVAL_THRESHOLD = slider_retrieval_threshold
    print("retrieval threshold:",model_settings.RETRIEVAL_THRESHOLD)

import yaml

def save_api_keys_to_yaml(GROQ_KEY, OPENAI_KEY, GEMINI_KEY):
    yaml_data = {
    "GROQ_API_KEY": GROQ_KEY,
    "OPENAI_API_KEY": OPENAI_KEY,
    "GEMINI_API_KEY": GEMINI_KEY,
    }
    with open('api_keys.yaml', 'w') as file:
        yaml.dump(yaml_data, file)

def load_api_keys_from_yaml():
    GROQ_KEY = ""
    OPENAI_KEY = ""
    GEMINI_KEY = ""

    if os.path.exists("api_keys.yaml"):
        with open("api_keys.yaml", 'r') as stream:
            data_loaded = yaml.safe_load(stream)
        
            GROQ_KEY = data_loaded["GROQ_API_KEY"]
            OPENAI_KEY = data_loaded["OPENAI_API_KEY"]
            GEMINI_KEY = data_loaded["GEMINI_API_KEY"]

            model_settings.GROQ_API_KEY = GROQ_KEY
            model_settings.OPENAI_API_KEY = OPENAI_KEY
            model_settings.GEMINI_API_KEY = GEMINI_KEY

            os.environ['GROQ_API_KEY'] = GROQ_KEY
            os.environ["OPENAI_API_KEY"] = OPENAI_KEY
            os.environ["GEMINI_API_KEY"] = GEMINI_KEY

    return GROQ_KEY, OPENAI_KEY, GEMINI_KEY

def btn_key_save_click(txt_groq_api_key, txt_openai_api_key, txt_gemini_api_key):
    model_settings.GROQ_API_KEY = txt_groq_api_key
    model_settings.OPENAI_API_KEY = txt_openai_api_key
    model_settings.GEMINI_API_KEY = txt_gemini_api_key

    os.environ['GROQ_API_KEY'] = txt_groq_api_key
    os.environ["OPENAI_API_KEY"] = txt_openai_api_key
    os.environ["GEMINI_API_KEY"] = txt_gemini_api_key

    save_api_keys_to_yaml(txt_groq_api_key, txt_openai_api_key, txt_gemini_api_key)

    print("\nSave API keys ~> Ok")

def get_groq_list_models(groq_api_key):
    url = "https://api.groq.com/openai/v1/models"
    headers = {
        "Authorization": f"Bearer {groq_api_key}",
        "Content-Type": "application/json"
    }
    response = requests.get(url, headers=headers)
    result = response.json()
    list_models = []
    for model in result['data']:
        list_models.append(model['id'])
    return list_models

from openai import OpenAI
def get_openai_list_models(openai_api_key):
    openai_client = OpenAI(api_key=openai_api_key)
    models = openai_client.models.list()
    list_models = []
    for model in models.data:
        list_models.append(model.id)
    return list_models

import google.generativeai as genai
def get_gemini_list_modes(gemini_api_key):
    genai.configure(api_key=gemini_api_key)
    list_models = []
    for m in genai.list_models():
      if 'generateContent' in m.supported_generation_methods:
        list_models.append(m.name.split("/")[-1])
    return list_models

# def get_litellm_list_models():
#     url = "http://0.0.0.0:4000/models"
#     headers = {
#         "accept": "application/json",
#     }
#     response = requests.get(url, headers=headers)
#     result = response.json()
#     return [result['data'][0]['id']]

def dropdown_model_type_select(dropdown_model_type):
    model_settings.MODEL_TYPE = dropdown_model_type
    print("\ndropdown_model_type:",model_settings.MODEL_TYPE)

def ollama_dropdown_model_select(dropdown_model):
    model_settings.MODEL_NAME = dropdown_model
    print("\nSelected model:",model_settings.MODEL_NAME)

def groq_dropdown_model_select(dropdown_model):
    model_settings.MODEL_NAME = dropdown_model
    print("\nSelected model:",model_settings.MODEL_NAME)

def openai_dropdown_model_select(dropdown_model):
    model_settings.MODEL_NAME = dropdown_model
    print("\nSelected model:",model_settings.MODEL_NAME)

def gemini_dropdown_model_select(dropdown_model):
    model_settings.MODEL_NAME = dropdown_model
    print("\nSelected model:",model_settings.MODEL_NAME)

def litellm_dropdown_model_select(dropdown_model):
    model_settings.MODEL_NAME = dropdown_model
    print("\nSelected model:",model_settings.MODEL_NAME)

def btn_create_new_workspace_click(workspace_list):
    max_id = 0
    for wp in workspace_list:
        if wp["id"] >= max_id:
            max_id = wp["id"] + 1
    workspace = {"id":max_id, "name":"New workspace "+str(max_id), "history":[["**human**: Hello", "**Jarvis (AI)**: Hi, my name Jarvis. I am your assistant. How may I help you today?  [v{0}]".format(max_id)]]}
    workspace_list.insert(0, workspace)
    return workspace_list, workspace

def update_is_retrieval(is_retrieval):
    model_settings.IS_RETRIEVAL = is_retrieval

def update_function_calling(function_calling):
    model_settings.FUNCTION_CALLING = function_calling

def btn_save_workspace_click(workspace_list):
    my_platform = platform.system() #  "Linux", "Windows", or "Darwin" (Mac)
    folder_path = "chat_workspaces"

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    # current date and time
    now = datetime.now()
    time_now = now.strftime("%Y-%m-%d_%H-%M-%S")

    for wp in workspace_list:
        file_name = str(time_now)+"_"+str(wp["id"])+"_"+str(wp["name"])+'.txt'

        file_path = ""
        if my_platform == "Windows":
            file_path = folder_path + "\\" + file_name
        elif my_platform == "Darwin":
           file_path = folder_path + "/" + file_name
        else:
            file_path = folder_path + "/" + file_name

        with open(file_path, 'w', encoding="utf-8") as f:
            for chat in wp["history"]:
                f.write(str(chat[0])+"\n"+str(chat[1])+"\n\n")
        print("\nsave workspace to ~>",file_path)


# Style class ------------------------------------------------------------
class UI_Style(Base):
    def __init__(
        self,
        *,
        primary_hue: colors.Color | str = colors.emerald,    # colors.cyan / colors.green / colors.yellow / colors.emerald
        secondary_hue: colors.Color | str = colors.blue,     # colors.blue / colors.orange / colors.red / colors.blue
        neutral_hue: colors.Color | str = colors.blue,       # colors.blue / colors.orange / colors.red / colors.gray
        spacing_size: sizes.Size | str = sizes.spacing_md,
        radius_size: sizes.Size | str = sizes.radius_md,
        text_size: sizes.Size | str = sizes.text_md,
        font: fonts.Font
        | str
        | Iterable[fonts.Font | str] = (
            fonts.GoogleFont("Quicksand"),
            "ui-sans-serif",
            "sans-serif",
        ),
        font_mono: fonts.Font
        | str
        | Iterable[fonts.Font | str] = (
            fonts.GoogleFont("IBM Plex Mono"),
            "ui-monospace",
            "monospace",
        ),
    ):
        super().__init__(
            primary_hue=primary_hue,
            secondary_hue=secondary_hue,
            neutral_hue=neutral_hue,
            spacing_size=spacing_size,
            radius_size=radius_size,
            text_size=text_size,
            font=font,
            font_mono=font_mono,
        )
        super().set(
            body_background_fill="repeating-linear-gradient(45deg, *primary_200, *primary_200 10px, *primary_50 10px, *primary_50 20px)",
            body_background_fill_dark="repeating-linear-gradient(45deg, *primary_800, *primary_800 10px, *primary_900 10px, *primary_900 20px)",
            
            button_primary_background_fill="linear-gradient(90deg, *primary_300, *secondary_400)",
            button_primary_background_fill_hover="linear-gradient(90deg, *primary_200, *secondary_300)",
            button_primary_text_color="white",
            button_primary_background_fill_dark="linear-gradient(90deg, *primary_600, *secondary_800)",
            
            slider_color="*secondary_300",
            slider_color_dark="*secondary_600",
            
            block_title_text_weight="600",
            block_border_width="3px",
            block_shadow="*shadow_drop_lg",
            
            button_shadow="*shadow_drop_lg",
            button_large_padding="12px",
        )

# theme_default = gr.themes.Default().set(
#     body_background_fill="repeating-linear-gradient(45deg, *primary_200, *primary_200 10px, *primary_50 10px, *primary_50 20px)",
#     body_background_fill_dark="repeating-linear-gradient(45deg, *primary_800, *primary_800 10px, *primary_900 10px, *primary_900 20px)",
# )

# theme_default_blue = gr.themes.Default(primary_hue="blue").set(
#     body_background_fill="repeating-linear-gradient(45deg, *primary_200, *primary_200 10px, *primary_50 10px, *primary_50 20px)",
#     body_background_fill_dark="repeating-linear-gradient(45deg, *primary_800, *primary_800 10px, *primary_900 10px, *primary_900 20px)",
# )


# GUI ------------------------------------------------------------
def JARVIS_assistant():
    ui_style = UI_Style()

    theme_1 = "gradio/default"
    # theme_2 = gr.Theme.from_hub("gradio/base")
    # theme_3 = gr.Theme.from_hub("gradio/seafoam")
    # theme_4 = gr.Theme.from_hub("gradio/glass")
    # theme_5 = gr.Theme.from_hub("gstaff/xkcd") 
    # theme_6 = gr.Theme.from_hub("ParityError/LimeFace")
    # theme_7 = gr.Theme.from_hub("EveryPizza/Cartoony-Gradio-Theme")
    # theme_8 = gr.Theme.from_hub("snehilsanyal/scikit-learn")
    # theme_9 = gr.Theme.from_hub("abidlabs/banana")
    
    with gr.Blocks(theme=ui_style) as GUI:
        with gr.Row():
            with gr.Column(scale=1):
                with gr.Tab("Workspace"):
                    first_state_workspace = {"id":0, "name":"My first workspace", "history":[["**human**: Hello", "**Jarvis (AI)**: Hi, my name Jarvis. I am your assistant. How may I help you today?"]]}
                    state_workspace_list = gr.State([first_state_workspace])
                    state_workspace_selected = gr.State(first_state_workspace)
    
                    with gr.Row(variant="panel"):
                        with gr.Row():
                            btn_save_workspace = gr.Button(value="Save all workspaces", min_width=220)
                            btn_save_workspace.click(fn=btn_save_workspace_click, inputs=state_workspace_list)
			    
                            btn_create_new_workspace = gr.Button(value="Create new workspace", min_width=220)
                            btn_create_new_workspace.click(fn=btn_create_new_workspace_click, inputs=state_workspace_list, outputs=[state_workspace_list, state_workspace_selected])
    
                        with gr.Row():
                            @gr.render(inputs=[state_workspace_list, state_workspace_selected])
                            def show_new_workspace(workspace_list, workspace_selected):
                                def txt_workspace_focus(focus_state):
                                    for wp in workspace_list:
                                        if wp["id"] == focus_state:
                                            return wp
                                
                                def btn_delete_workspace_click(focus_state):
                                    if len(workspace_list) > 1:
                                        for wp in workspace_list:
                                            if wp["id"] == focus_state:
                                                workspace_list.remove(wp)
                                    return workspace_list
                                
                                def txt_workspace_change(focus_state, txt_workspace):
                                    for wp in workspace_list:
                                        if wp["id"] == focus_state:
                                            workspace_list.remove(wp)
                                            workspace= {"id":wp["id"], "name":txt_workspace, "history":wp["history"]}
                                            workspace_list.insert(0, workspace)
                                            return workspace_list, workspace
    
                                for wksp in workspace_list:
                                    with gr.Row():
                                        with gr.Column(scale=10, min_width=200):
                                            focus_state = gr.State(wksp["id"])
                                            txt_workspace = gr.Textbox(value=str(wksp["name"]), show_label=False, container=False, min_width=200,  interactive=True)
                                            txt_workspace.focus(txt_workspace_focus, focus_state, [state_workspace_selected])
                                            txt_workspace.submit(txt_workspace_change, [focus_state, txt_workspace], [state_workspace_list, state_workspace_selected])
                                        
                                        with gr.Column(scale=1, min_width=10):
                                            # if wksp["id"] != 0:
                                            btn_delete_workspace = gr.Button(value="x", min_width=5) #  size="sm"
                                            btn_delete_workspace.click(fn=btn_delete_workspace_click, inputs=focus_state, outputs=[state_workspace_list])
                
                with gr.Tab("Model"):
                    with gr.Row():
                        with gr.Row(variant="panel"):
                            with gr.Accordion(label="API Keys", open=False):
                                GROQ_KEY, OPENAI_KEY, GEMINI_KEY = load_api_keys_from_yaml()
                                txt_groq_api_key = gr.Textbox(value=GROQ_KEY, placeholder="GroqCloud API Key", show_label=False)
                                txt_openai_api_key = gr.Textbox(value=OPENAI_KEY, placeholder="OpenAI API Key", show_label=False)
                                txt_gemini_api_key = gr.Textbox(value=GEMINI_KEY, placeholder="Gemini API Key", show_label=False)
                                
                                btn_key_save = gr.Button(value="Save", min_width=50)
                                btn_key_save.click(fn=btn_key_save_click, inputs=[txt_groq_api_key, txt_openai_api_key, txt_gemini_api_key])
                                
                        with gr.Row(variant="panel"):
    
                            dropdown_model_type = gr.Dropdown(choices=["Ollama", "GroqCloud", "OpenAI", "Gemini", "LiteLLM"], value=model_settings.MODEL_TYPE, type="value", label="Type", interactive=True, min_width=220)
                            dropdown_model_type.select(fn=dropdown_model_type_select, inputs=[dropdown_model_type])
    
                            @gr.render(inputs=dropdown_model_type)
                            def show_dropdown_model(dropdown_model_type):
                                if dropdown_model_type == "Ollama":
                                    ollama_list_models = get_ollama_list_models()
                                    model_settings.MODEL_NAME = ollama_list_models[0]
                                    print("Selected model:",model_settings.MODEL_NAME)
    
                                    ollama_dropdown_model = gr.Dropdown(choices=ollama_list_models, value=model_settings.MODEL_NAME, type="value", label="Model", interactive=True)
                                    ollama_dropdown_model.select(fn=ollama_dropdown_model_select, inputs=[ollama_dropdown_model])
    
                                if dropdown_model_type == "GroqCloud" and model_settings.GROQ_API_KEY:
                                    groq_list_models = get_groq_list_models(model_settings.GROQ_API_KEY)
                                    model_settings.MODEL_NAME = groq_list_models[0]
                                    print("Selected model:",model_settings.MODEL_NAME)
    
                                    groq_dropdown_model = gr.Dropdown(choices=groq_list_models, value=model_settings.MODEL_NAME, type="value", label="Models", interactive=True)
                                    groq_dropdown_model.select(fn=groq_dropdown_model_select, inputs=[groq_dropdown_model])
    
                                if dropdown_model_type == "OpenAI" and model_settings.OPENAI_API_KEY:
                                    openai_list_models = get_openai_list_models(model_settings.OPENAI_API_KEY)
                                    model_settings.MODEL_NAME = openai_list_models[0]
                                    print("Selected model:",model_settings.MODEL_NAME)
    
                                    openai_dropdown_model = gr.Dropdown(choices=openai_list_models, value=model_settings.MODEL_NAME, type="value", label="Models", interactive=True)
                                    openai_dropdown_model.select(fn=openai_dropdown_model_select, inputs=[openai_dropdown_model])
    
                                if dropdown_model_type == "Gemini" and model_settings.GEMINI_API_KEY:
                                    gemini_list_modes = get_gemini_list_modes(model_settings.GEMINI_API_KEY)
                                    model_settings.MODEL_NAME = gemini_list_modes[0]
                                    print("Selected model:",model_settings.MODEL_NAME)
    
                                    gemini_dropdown_model = gr.Dropdown(choices=gemini_list_modes, value=model_settings.MODEL_NAME, type="value", label="Models", interactive=True)
                                    gemini_dropdown_model.select(fn=gemini_dropdown_model_select, inputs=[gemini_dropdown_model])
    
                                if dropdown_model_type == "LiteLLM":
                                    litellm_list_models = get_ollama_list_models()
                                    model_settings.MODEL_NAME = litellm_list_models[0]
                                    print("Selected model:",model_settings.MODEL_NAME)
    
                                    litellm_dropdown_model = gr.Dropdown(choices=litellm_list_models, value=model_settings.MODEL_NAME, type="value", label="Models", interactive=True)
                                    litellm_dropdown_model.select(fn=litellm_dropdown_model_select, inputs=[litellm_dropdown_model])
    
                            # radio_device = gr.Radio(choices=["GPU", "MLX", "CPU"], value='CPU', label="Device")
                            # radio_device.select(fn=radio_device_select, inputs=[radio_device])
    
                        with gr.Row(variant="panel"):
                            with gr.Accordion(label="Model settings", open=True):
                                slider_num_predict = gr.Slider(minimum=0, maximum=4096, value=model_settings.NUM_PREDICT, step=256, label="Max new tokens", interactive=True, min_width=220)
                                slider_num_predict.change(fn=slider_num_predict_change, inputs=slider_num_predict)
    
                                slider_temperature = gr.Slider(minimum=0, maximum=1, value=model_settings.TEMPERATURE, step=0.1, label="Temperature", interactive=True)
                                slider_temperature.change(fn=slider_temperature_change, inputs=slider_temperature)
    
                                slider_top_k = gr.Slider(minimum=0, maximum=100, value=model_settings.TOP_K, step=10, label="Top_k", interactive=True)
                                slider_top_k.change(fn=slider_top_k_change, inputs=slider_top_k)
    
                                slider_top_p = gr.Slider(minimum=0, maximum=1, value=model_settings.TOP_P, step=0.05, label="Top_p", interactive=True)
                                slider_top_p.change(fn=slider_top_p_change, inputs=slider_top_p)

                                chk_function_calling = Toggle(label="Function calling", value=True, interactive=True)
                                chk_function_calling.change(fn=update_function_calling, inputs=chk_function_calling)
    
                        with gr.Row(variant="panel"):
                            with gr.Accordion(label="Retrieval settings", open=True):
                                chk_is_retrieval = Toggle(label="Is retrieval", value=True, interactive=True)
                                chk_is_retrieval.change(fn=update_is_retrieval, inputs=chk_is_retrieval)
                                
                                slider_retrieval_top_k = gr.Slider(minimum=1, maximum=30, value=model_settings.RETRIEVAL_TOP_K, step=1, label="Top-K", interactive=True, min_width=220)
                                slider_retrieval_top_k.change(fn=slider_retrieval_top_k_change, inputs=slider_retrieval_top_k)
    
                                slider_retrieval_threshold = gr.Slider(minimum=0, maximum=1, value=model_settings.RETRIEVAL_THRESHOLD, step=0.05, label="Threshold score", interactive=True)
                                slider_retrieval_threshold.change(fn=slider_retrieval_threshold_change, inputs=slider_retrieval_threshold)
    
                with gr.Tab("System prompt"):
                    with gr.Row():
                        txt_system_prompt = gr.Textbox(value=system_prompt, label="System prompt", lines=23, min_width=220)
    
                        with gr.Row():
                            with gr.Column(scale=1, min_width=50):
                                btn_save = gr.Button(value="Save")
                                btn_save.click(fn=btn_save_click, inputs=[txt_system_prompt])
    
                            with gr.Column(scale=1, min_width=50):
                                btn_reset = gr.Button(value="Reset")
                                btn_reset.click(fn=btn_reset_click, inputs=txt_system_prompt, outputs=txt_system_prompt)
                
            with gr.Column(scale=7):
                def update_chat_history(chatbot, workspace_list, workspace_selected):
                    for wp in workspace_list:
                        if wp["id"] == workspace_selected["id"]:
                            workspace= {"id":wp["id"], "name":wp["name"], "history":chatbot}
                            workspace_list.remove(wp)
                            workspace_list.insert(0, workspace)
                            return workspace_list, workspace
    
                workspace_selected = state_workspace_selected.value
                chatbot = gr.Chatbot(workspace_selected["history"], elem_id="chatbot", bubble_full_width=False, min_width=800, height=560, show_copy_button=True,)
                chat_input = gr.MultimodalTextbox(value={"text": ""}, interactive=True, file_types=[".pdf",".txt"], file_count='multiple', placeholder="Enter message or upload file...", show_label=False)
    
                def workspace_selected_chatbot(workspace_selected):
                    return workspace_selected["history"]
                state_workspace_selected.change(fn=workspace_selected_chatbot, inputs=state_workspace_selected,  outputs=chatbot)
    
                chat_msg = chat_input.submit(fn=add_message, inputs=[chatbot, chat_input], outputs=[chatbot])
                bot_msg = chat_msg.then(fn=bot, inputs=[chatbot, chat_input], outputs=[chatbot, chat_input]).then(fn=update_chat_history, inputs=[chatbot, state_workspace_list, state_workspace_selected], outputs=[state_workspace_list, state_workspace_selected])
    chatbot
    
                # gr.Examples(examples=[{'text': "Bạn tên là gì?"}, {'text': "What's your name?"}, {'text': 'Quel est ton nom?'}, {'text': 'Wie heißen Sie?'}, {'text': '¿Cómo te llamas?'}, {'text': '你叫什么名字？'}, {'text': 'あなたの名前は何ですか？'}, {'text': '이름이 뭐에요?'}, {'text': 'คุณชื่ออะไร?'}, {'text': 'ما اسمك؟'}], inputs=chat_input)
    
    GUI.launch()

if __name__ == "__main__":
    JARVIS_assistant()
